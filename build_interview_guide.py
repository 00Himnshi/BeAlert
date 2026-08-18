from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Assignment_Alerts_Interview_Guide.docx"
BLUE = "2E74B5"
NAVY = "1F4D78"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F9"
MUTED = RGBColor(89, 89, 89)


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_text(p, text, bold=False, italic=False, color=None, size=11):
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return r


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    add_text(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    add_text(p, text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    add_text(p, text)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_text(p, text, bold=True, color=BLUE if level < 3 else NAVY, size={1: 16, 2: 13, 3: 12}[level])
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, f"{label}: ", bold=True, color=NAVY)
    add_text(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table_geometry(table, [2700, 6660])
    for label, value in rows:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        add_text(p, label, bold=True)
        p = cells[1].paragraphs[0]
        add_text(p, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_process_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table_geometry(table, [1900, 3500, 3960])
    for cell, value in zip(table.rows[0].cells, ("Stage", "What happens", "Why it matters")):
        shade(cell, BLUE)
        p = cell.paragraphs[0]
        add_text(p, value, bold=True, color="FFFFFF")
    for stage, action, why in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (stage, action, why)):
            add_text(cell.paragraphs[0], text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
for level, size, before, after, color in [(1, 16, 16, 8, BLUE), (2, 13, 11, 5, BLUE), (3, 12, 7, 3, NAVY)]:
    st = styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_text(header, "Assignment Alerts | Interview Guide", size=9, color="666666")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(footer, "Personal project technical walkthrough", size=9, color="666666")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(5)
add_text(p, "TECHNICAL WALKTHROUGH", size=10, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
add_text(p, "Assignment Alerts", size=25, bold=True, color="1F1F1F")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(15)
add_text(p, "How the portal checker, Supabase database, React dashboard, Twilio alerts, and GitHub Actions work together", size=13, color="555555")
add_key_value_table(doc, [
    ("Purpose", "Automatically discover new assignments from a Moodle-based university portal and notify the student."),
    ("Main idea", "Turn repeated manual portal checking into a scheduled, idempotent background job."),
    ("Schedule", "GitHub Actions runs at 00:07, 08:07, and 16:07 UTC (every 8 hours)."),
    ("Core stack", "Python, Requests, Beautiful Soup, Supabase/Postgres, React/Vite, Twilio WhatsApp, GitHub Actions."),
])
add_callout(doc, "Interview one-liner", "I built an end-to-end assignment-monitoring workflow. Its key design problem was reliable change detection: I assign each portal activity a stable ID, compare it to the database, and alert only when an ID is genuinely new.")

add_heading(doc, "1. System overview")
add_body(doc, "The project has two separate paths. The private backend path checks the university portal, stores data, and sends alerts. The public frontend path is a React dashboard that reads saved data after the owner signs in. Keeping these paths separate means browser users never receive portal credentials or Supabase's server-side key.")
add_process_table(doc, [
    ("Scheduled trigger", "GitHub Actions starts the Python script every eight hours, or when started manually.", "No personal computer needs to stay on."),
    ("Portal check", "The script logs in, downloads the configured course page, and parses assignment-looking activities.", "Produces the current list of assignments."),
    ("Database comparison", "The script reads known portal IDs from Supabase and compares them with the IDs just found.", "Separates new assignments from already-known ones."),
    ("Notification", "One Twilio WhatsApp template message is sent if one or more IDs are new.", "Avoids repeat messages and reduces notification noise."),
    ("Dashboard", "React fetches assignments from Supabase after the configured user authenticates.", "Shows the same saved source of truth in a usable interface."),
])

add_heading(doc, "2. How Supabase stores assignments")
add_body(doc, "Supabase provides a hosted Postgres database. The table is created by supabase/database.sql. Each row represents one assignment activity found on the course page.")
add_key_value_table(doc, [
    ("id", "A database-generated UUID primary key. Useful as an internal row identifier."),
    ("portal_id", "Text identifier for the Moodle activity. It is NOT NULL and UNIQUE, which is the main duplicate-prevention rule."),
    ("title", "The assignment title extracted from the course page."),
    ("assignment_url", "Direct link to the assignment activity."),
    ("course_url", "The course page that was checked. This allows the data to retain its source."),
    ("first_seen_at", "Timestamp set when the row is first inserted."),
    ("last_seen_at", "Timestamp updated each time the assignment is seen in a later scan."),
])
add_callout(doc, "Important point", "The UNIQUE constraint on portal_id is a database-level safety net. Even if the script is accidentally run twice, Supabase will not create two rows for the same assignment ID.")

add_heading(doc, "3. Creating a stable assignment ID")
add_body(doc, "A title is not reliable enough to identify an assignment because a teacher may rename it. The checker instead tries to use Moodle's activity ID from the assignment URL. For example, a URL containing ?id=123 becomes the stable ID moodle-123.")
add_body(doc, "If an activity URL has no id query parameter, the code generates a SHA-256 hash of the full URL. That creates a repeatable fallback identifier: the same URL always produces the same hash.")
add_key_value_table(doc, [
    ("Preferred identifier", "moodle-<activity-id>, extracted from the URL query string."),
    ("Fallback identifier", "SHA-256 hash of the assignment URL."),
    ("Why not use the title?", "Titles can be edited; URLs/activity IDs are more stable for identifying the same resource."),
])

add_heading(doc, "4. How old and current assignments are matched")
add_body(doc, "This is the central change-detection flow. On every run, the checker creates a current list from the portal and asks Supabase for the set of portal_id values that it already knows. It then uses a set-membership comparison.")
add_process_table(doc, [
    ("1. Read portal", "find_assignments() parses li.activity a[href] links and filters titles containing assignment or lab exercise.", "Builds Assignment objects with portal_id, title, and assignment_url."),
    ("2. Read old IDs", "existing_portal_ids() calls the Supabase REST API with select=portal_id.", "Only the small identifier set is downloaded, not every field."),
    ("3. Compare", "new_assignments = [item for item in found_assignments if item.portal_id not in known_ids]", "An assignment is new only when its stable ID is absent from the old-ID set."),
    ("4. Save", "save_assignments() upserts all currently found rows using on_conflict=portal_id.", "New rows are inserted; existing rows are updated, including last_seen_at."),
])
add_callout(doc, "Plain-English explanation", "The database acts like memory. The portal tells us what exists now; Supabase tells us what the checker has seen before. Anything that exists now but is not in that memory is treated as new.")

add_heading(doc, "5. First run, later runs, and retry behavior")
add_body(doc, "The first run needs special handling. If the database has no known IDs yet, every assignment on the portal would technically look new. Instead, the script saves them as a baseline and deliberately sends no WhatsApp message. This gives future scans a starting point.")
add_process_table(doc, [
    ("First successful run", "No known IDs exist. Save all currently found assignments; do not alert.", "Prevents old course work from generating a large, misleading notification."),
    ("Later run: no change", "No current ID is missing from known_ids. Upsert current rows and print 'No new assignments found.'", "Keeps last_seen_at current without duplicate alerts."),
    ("Later run: new item", "Send one WhatsApp message for all new assignments, then save/upsert rows.", "The student gets one concise alert."),
    ("Twilio fails", "The new rows are not saved because saving occurs after sending the message.", "The next scheduled run can retry the alert instead of silently losing it."),
])

add_heading(doc, "6. How the database write works")
add_body(doc, "The checker uses Supabase's REST endpoint rather than a direct database connection. It POSTs a JSON array of rows to /rest/v1/assignments?on_conflict=portal_id with the header Prefer: resolution=merge-duplicates.")
add_body(doc, "This is an upsert operation: if portal_id does not exist, Supabase inserts a new row and sets first_seen_at automatically. If portal_id already exists, Supabase merges the new data into the existing row. In this project, last_seen_at is explicitly refreshed to the current UTC time on every successful scan.")
add_callout(doc, "Why upsert?", "It makes repeated scans safe. A checker can see the same assignment many times without producing duplicate database rows, while still recording that the assignment was visible in the latest scan.")

add_heading(doc, "7. Security model")
add_body(doc, "The automation uses credentials, so the project separates secret server operations from browser operations.")
add_process_table(doc, [
    ("GitHub Secrets", "Stores portal username/password, Twilio credentials, and the Supabase server-side key.", "Secrets are injected as environment variables only during the workflow run."),
    ("Server-side Supabase key", "Used only by the Python checker to read and write assignments.", "It can bypass Row Level Security, so it must never enter React code."),
    ("Supabase anon key", "Used by the dashboard build as a GitHub Variable.", "It is safe for the browser, but has limited permissions."),
    ("Row Level Security", "The SQL policy allows SELECT only for the configured authenticated email address.", "The dashboard user can read their assignments but cannot publicly browse or write the table."),
])
add_body(doc, "One limitation to state honestly: the configured university portal login endpoint uses HTTP. That is a property of the portal, not an encryption choice made by this project. A dedicated low-privilege account is safer than reusing a high-value password.")

add_heading(doc, "8. How React renders updates on the dashboard")
add_body(doc, "The React application is a separate Vite project in dashboard/. It does not scrape the portal and it does not calculate changes. Its job is to display the data that the checker has already saved to Supabase.")
add_process_table(doc, [
    ("Configuration", "App.jsx reads VITE_SUPABASE_URL and VITE_SUPABASE_ANON_KEY from the built frontend environment.", "Creates the Supabase browser client without exposing server secrets."),
    ("Authentication", "The user signs in with Supabase email/password authentication.", "The database policy can identify the user from their authenticated JWT."),
    ("Query", "loadAssignments() selects title, assignment_url, first_seen_at, and last_seen_at from assignments, ordered by first_seen_at descending.", "Fetches the latest saved list from the single source of truth."),
    ("State update", "The returned rows are stored in React state with setAssignments(data).", "Changing state triggers React to render the new list."),
    ("Display", "assignments.map(...) creates one article per row with title, link, and first-found timestamp.", "New assignments appear the next time the list is loaded or the Refresh list button is clicked."),
])
add_callout(doc, "Important distinction", "The current dashboard uses fetch-on-load plus a Refresh list button. It is not using Supabase Realtime subscriptions. If real-time updates were needed, I would add a channel subscription or periodic client-side refresh.")

add_heading(doc, "9. How GitHub Actions runs the automation")
add_body(doc, "The workflow file .github/workflows/check-assignments.yml defines two triggers: schedule and workflow_dispatch. schedule starts the checker every eight hours at 00:07, 08:07, and 16:07 UTC; workflow_dispatch adds a Run workflow button for manual testing.")
add_process_table(doc, [
    ("Checkout", "actions/checkout@v4 downloads the repository into the GitHub-hosted runner.", "The runner needs the checker source and requirements file."),
    ("Python setup", "actions/setup-python@v5 installs Python 3.12.", "Makes the runtime predictable."),
    ("Dependencies", "pip install -r checker/requirements.txt installs Requests, Beautiful Soup, Twilio, and dotenv.", "Gives the script its required libraries."),
    ("Run script", "python checker/check_assignments.py executes the workflow logic.", "The env section maps GitHub Secrets to the variable names expected by the code."),
])
add_body(doc, "The dashboard has a separate publish-dashboard.yml workflow. It runs on pushes to main or manually, installs Node.js dependencies, builds the Vite application with the Supabase browser variables, and deploys dashboard/dist to GitHub Pages.")

add_heading(doc, "10. Strong interview answers")
add_heading(doc, "What problem did you solve?", 2)
add_body(doc, "I automated a repetitive academic task: checking a course portal for new assignments. The system runs independently of my laptop, remembers previous results, sends a WhatsApp alert only for new work, and gives me a private dashboard to review what it found.")
add_heading(doc, "How did you prevent duplicate alerts?", 2)
add_body(doc, "I gave every assignment a stable portal_id based on Moodle's activity ID, with a URL-hash fallback. Before alerting, the script loads known IDs from Supabase and alerts only for IDs not present in that set. The database also has a UNIQUE constraint on portal_id as a second layer of protection.")
add_heading(doc, "Why did you use an upsert?", 2)
add_body(doc, "A scheduled job sees the same assignments repeatedly. Upsert makes the write idempotent: it inserts a new assignment once but safely refreshes last_seen_at for existing assignments. That removes the need for separate insert-versus-update handling.")
add_heading(doc, "How is the dashboard secured?", 2)
add_body(doc, "The dashboard uses Supabase's anon key and authenticated user session, never the service-role key. Row Level Security allows reads only for the configured email, and there is no browser-facing insert or update policy.")
add_heading(doc, "What would you improve next?", 2)
add_bullet(doc, "Support multiple course URLs and store course configuration in the database.")
add_bullet(doc, "Use Moodle APIs or stronger selectors if available, because HTML scraping depends on page structure.")
add_bullet(doc, "Add automated tests using saved sample HTML pages and mock Supabase/Twilio responses.")
add_bullet(doc, "Add observability: failed-run alerts, structured logs, and a status/history table.")
add_bullet(doc, "Add Supabase Realtime subscriptions if the dashboard should update without pressing Refresh.")

add_heading(doc, "11. Final 30-second explanation")
add_callout(doc, "Suggested answer", "Assignment Alerts is a scheduled full-stack automation I built for a Moodle portal. GitHub Actions runs a Python checker every eight hours. The checker logs in, extracts assignment activities, and compares their stable IDs with IDs stored in Supabase. Only missing IDs are treated as new, so the project avoids duplicate notifications. New items trigger one Twilio WhatsApp alert and are upserted into Postgres. A React dashboard then reads the same saved records through Supabase authentication and Row Level Security. The main engineering focus was making the scheduled job safe to run repeatedly.")

doc.save(OUT)
print(OUT)
